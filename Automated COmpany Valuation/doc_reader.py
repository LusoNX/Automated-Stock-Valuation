from __future__ import print_function
from mailmerge import MailMerge
from datetime import date
import yfinance as yf 
import sklearn
from sklearn.linear_model import LinearRegression
import numpy as np 
from pandas import json_normalize
import pandas as pd 

from docx import Document

import Yahoo_Valuation as yv
import DCF_VALUATION as DCF

template = "company-presentation-template.docx"

document = MailMerge(template)

df_industry = pd.read_csv("industry_multiples_FINAL.csv")
df_sector = pd.read_csv("sector_multiples_FINAL.csv")

def document_reader(symbol):
	_ticker = yf.Ticker(symbol)
	sym = _ticker.info
	try:
		_description = sym["longBusinessSummary"]
	except:
		print("Stock ticker does not exist.")
	currency = sym["financialCurrency"]
	ke = yv.cost_of_equity(symbol) 
	ke = np.round(ke,2)
	wacc = yv.wacc(symbol)
	wacc = np.round(wacc,2)

	ke = str(ke*100) + "%"
	wacc = str(wacc*100) + "%"

	pprice = np.round(sym["currentPrice"],2)
	_price = str(np.round(sym["currentPrice"],2)) +" " + currency
	_name = sym["shortName"]



	_industry = sym["industry"]
	_sector = sym["sector"]
	_country = sym["country"]


	## Company Ratios
	gross_margin_company = str(np.round(sym["grossMargins"]*100,1)) + "%"
	operating_margin_company = str(np.round(sym["operatingMargins"]*100,1)) + "%"
	ni_margin_company = str(np.round(sym["profitMargins"]*100,1))+ "%"
	ROA_company = str(np.round(sym["returnOnAssets"]*100,1))+ "%"
	if sym["returnOnEquity"] != None:
		ROE_company = str(np.round(sym["returnOnEquity"]*100,1))+ "%"

	else:
		ROE_company = "N/A"


	if sym["debtToEquity"] != None:
		debt_to_equity_company = str(np.round(sym["debtToEquity"],1))+ "%"

	else:
		debt_to_equity_company = "N/A"

	try: 
		PE_company = str(np.round(sym["trailingPE"],2))
	except KeyError:
		PE_company = "N/A"
		pass

	
	if sym["priceToBook"] != None:
		PB_company = str(np.round(sym["priceToBook"],2))

	else:
		PB_company = "N/A"
		
	
	if sym["enterpriseToEbitda"] != None:
		EV_to_ebitda_company = str(sym["enterpriseToEbitda"])

	else:
		EV_to_ebitda_company = "N/A"

	# Industry ratios
	industry_values = df_industry[df_industry["industry"] == _industry]
	gross_margin_industry = str(np.round(industry_values["grossMargins"].values[0]*100,1)) + "%"
	operating_margin_industry = str(np.round(industry_values["operatingMargins"].values[0]*100,1)) + "%"	
	ni_margin_industry = str(np.round(industry_values["profitMargins"].values[0]*100,1)) +"%"
	ROA_industry = str(np.round(industry_values["returnOnAssets"].values[0]*100,1))+ "%"
	ROE_industry = str(np.round(industry_values["returnOnEquity"].values[0]*100,1))+ "%"
	debt_to_equity_industry = str(np.round(industry_values["debtToEquity"].values[0],1))+ "%"
	PE_industry = str(np.round(industry_values["trailingPE"].values[0],1))
	PB_industry = str(np.round(industry_values["priceToBook"].values[0],1))
	EV_to_ebitda_industry = str(np.round(industry_values["enterpriseToEbitda"].values[0],1))


	# Sector ratios
	sector_values = df_sector[df_sector["sector"] == _sector]
	gross_margin_sector = str(np.round(sector_values["grossMargins"].values[0]*100,1)) + "%"
	operating_margin_sector = str(np.round(sector_values["operatingMargins"].values[0]*100,1)) + "%"	
	ni_margin_sector = str(np.round(sector_values["profitMargins"].values[0]*100,1)) +"%"
	ROA_sector = str(np.round(sector_values["returnOnAssets"].values[0]*100,1))+ "%"
	ROE_sector = str(np.round(sector_values["returnOnEquity"].values[0]*100,1))+ "%"
	debt_to_equity_sector = str(np.round(sector_values["debtToEquity"].values[0],1))+ "%"
	PE_sector = str(np.round(sector_values["trailingPE"].values[0],1))
	PB_sector = str(np.round(sector_values["priceToBook"].values[0],1))
	EV_to_ebitda_sector = str(np.round(sector_values["enterpriseToEbitda"].values[0],1))

	#

	document.merge(ticker = symbol)
	document.merge(name = _name)
	document.merge(price = _price)
	document.merge(description = _description)
	document.merge(industry = _industry)
	document.merge(sector = _sector)
	document.merge(country = _country)


	## Analyst Consensus

	_analyst_consensus = str(sym["targetMeanPrice"])
	_nr_analyst = str(sym["numberOfAnalystOpinions"])

	# COmpany ratios DOC
	document.merge(gross_company = gross_margin_company)
	document.merge(operating_company = operating_margin_company)
	document.merge(ni_company = ni_margin_company)
	document.merge(roa_company = ROA_company)
	document.merge(roe_company = ROE_company)
	document.merge(de_company = debt_to_equity_company)
	document.merge(pe_company = PE_company)
	document.merge(pb_company = PB_company)
	document.merge(ev_to_ebitda_company = EV_to_ebitda_company)

	# Industry Ratios DOC
	document.merge(gross_industry = gross_margin_industry)
	document.merge(operating_industry = operating_margin_industry)
	document.merge(ni_company = ni_margin_industry)
	document.merge(roa_industry = ROA_industry)
	document.merge(roe_industry = ROE_industry)
	document.merge(de_industry = debt_to_equity_industry)
	document.merge(pe_industry = PE_industry)
	document.merge(pb_industry = PB_industry)
	document.merge(ev_to_ebitda_industry = EV_to_ebitda_industry)

	# Industry Ratios DOC
	document.merge(gross_sector = gross_margin_sector)
	document.merge(operating_sector = operating_margin_sector)
	document.merge(ni_sector = ni_margin_sector)
	document.merge(roa_sector = ROA_sector)
	document.merge(roe_sector = ROE_sector)
	document.merge(de_sector = debt_to_equity_sector)
	document.merge(pe_sector = PE_sector)
	document.merge(pb_sector = PB_sector)
	document.merge(ev_to_ebitda_sector = EV_to_ebitda_sector)


	document.merge(analyst_consensus=_analyst_consensus)
	document.merge(nr_analyst = _nr_analyst)

	document.merge(cost_of_equity = ke)
	document.merge(wacc = wacc)


	# Model valuation 
	# (1) Monte Carlo simulation 

	try:
		perc_25,_mean_value = DCF.dcf_valuation(symbol)
	except ValueError:
		perc_25,_mean_value = [0,0]
	document.merge(P25 = str(round(perc_25,2)))
	document.merge(mean_value = str(round(_mean_value,2)))


	# (2) Yahoo Finance FCF	Model
	estimate_A = yv.dcf_valuation_A(symbol,"A1")
	return_A = (pprice/estimate_A) -1
	

	

	estimate_B = yv.dcf_valuation_A(symbol,"A2")
	estimate_C = yv.dcf_valuation_A(symbol,"A3")

	# (3) Yahoo Finance Div Model

	#sector_values = df_sector[df_sector["industry"] == _sector]

	estimate_D = yv.div_valuation_B(symbol,"B1")
	estimate_E = yv.div_valuation_B(symbol,"B2")


	
	estimates = [estimate_A,estimate_B,estimate_C,estimate_D,estimate_E]


	target_returns = []
	for i in estimates:
		if i > 0:
			x_return = (i/pprice) -1
			target_returns.append(x_return)
		else:
			target_returns.append(0)


	estimates = np.array(estimates)
	estimates[estimates ==0] = np.nan
	mean_estimate = np.nanmean(estimates,axis =0)
	mean_estimate = str(np.round(mean_estimate,2))

	document.merge(target_price = str(round(_mean_value,2)))

	document.merge(target_price_a1 = str(round(estimate_A,2)))
	document.merge(target_price_a2 = str(round(estimate_B,2)))
	document.merge(target_price_a3 = str(round(estimate_C,2)))
	document.merge(target_price_b1 = str(round(estimate_D,2)))
	document.merge(target_price_b2 = str(round(estimate_E,2)))

	document.merge(potential_ret_a = str(round(target_returns[0],2)*100)+ "%") 
	document.merge(potential_ret_b = str(round(target_returns[1],2)*100)+ "%") 
	document.merge(potential_ret_c = str(round(target_returns[2],2)*100)+ "%") 
	document.merge(potential_ret_d = str(round(target_returns[3],2)*100)+ "%") 
	document.merge(potential_ret_e = str(round(target_returns[4],2)*100)+ "%") 
	document.write('test-output.docx')

	## THis second part is for the adding of immages to the document
	doc = Document("test-output.docx")

	image_paras = [i for i, p in enumerate(doc.paragraphs) if "[monte_carlo]" in p.text] ## Use this to identify the paragraph of a specific wording
	p = doc.paragraphs[image_paras[0]]
	p.text = ""
	r = p.add_run()
	try:
		r.add_picture("{i}_monte_carlo.png".format(i=symbol))
	except FileNotFoundError:
		pass
	doc.save("{i} Summary.docx".format(i = symbol))






document_reader("OHI")
